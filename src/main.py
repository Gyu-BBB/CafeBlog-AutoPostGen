# 이미지 및 텍스트 분석을 통한 자동 게시글 생성 시스템
# /src/main.py
import os
import sys
import urllib.request
import urllib.parse
import json
import re
import openai
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from transformers import BlipProcessor, BlipForConditionalGeneration
from PIL import Image
import torch

def get_api_keys():
    # API 키 및 클라이언트 정보 가져오기
    api_key = os.getenv("OPENAI_API_KEY")
    client_id = os.getenv("NAVER_CLIENT_ID")
    client_secret = os.getenv("NAVER_CLIENT_SECRET")
    
    if not api_key or not client_id or not client_secret:
        print("API 키 또는 클라이언트 정보가 설정되지 않았습니다.")
        sys.exit()
    
    return api_key, client_id, client_secret

def create_openai_client(api_key):
    # OpenAI API 키 설정
    openai.api_key = api_key

def convert_images_to_png(folder_path):
    """
    지정된 폴더 내의 모든 .jpg, .jpeg, .png 이미지를 .png 형식으로 변환합니다.
    변환된 이미지는 동일한 이름으로 .png 확장자를 가집니다.
    """
    supported_extensions = (".jpg", ".jpeg", ".png")
    for filename in os.listdir(folder_path):
        if filename.lower().endswith(supported_extensions):
            image_path = os.path.join(folder_path, filename)
            try:
                with Image.open(image_path) as img:
                    # 변환된 파일 이름 설정
                    png_filename = os.path.splitext(filename)[0] + ".png"
                    png_path = os.path.join(folder_path, png_filename)
                    
                    # 이미 .png 형식이면 건너뜀
                    if filename.lower().endswith(".png"):
                        continue
                    
                    # 이미지 변환 및 저장
                    img.convert("RGBA").save(png_path, "PNG")
                    print(f"이미지 변환 완료: {png_filename}")
            except Exception as e:
                print(f"이미지 변환 실패: {filename}, 오류: {e}")

def analyze_images_in_folder():
    folder_path = "/data/test/"
    
    processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base")
    model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base")

    captions = []
    image_filenames = []

    for filename in os.listdir(folder_path):
        if filename.lower().endswith(".png"):  # 변환된 .png 파일만 처리
            image_path = os.path.join(folder_path, filename)
            try:
                image = Image.open(image_path)
                inputs = processor(images=image, return_tensors="pt")
                out = model.generate(**inputs, max_new_tokens=50)
                caption = processor.decode(out[0], skip_special_tokens=True)
                captions.append(f"{caption} {{{filename}}}")
                image_filenames.append(filename)
                print(f"파일: {filename}, 생성된 캡션: {caption}")
            except Exception as e:
                print(f"이미지 분석 실패: {filename}, 오류: {e}")

    return captions, image_filenames

def read_sys_prompt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        sys_prompt = json.load(file)
    return sys_prompt

def generate_keywords(first_sys_prompt_content, user_question):
    completion = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": first_sys_prompt_content
            },
            {
                "role": "user",
                "content": user_question
            }
        ]
    )

    keyword = completion.choices[0].message.content.strip()
    print(f"추출된 키워드: {keyword}")
    return keyword

def search_naver_blog(client_id, client_secret, keyword):
    encText = urllib.parse.quote(keyword)
    display = 5
    start = 1
    sort = "sim"
    
    url = f"https://openapi.naver.com/v1/search/blog?query={encText}&display={display}&start={start}&sort={sort}"
    
    request = urllib.request.Request(url)
    request.add_header("X-Naver-Client-Id", client_id)
    request.add_header("X-Naver-Client-Secret", client_secret)
    
    try:
        response = urllib.request.urlopen(request)
        response_body = response.read()
        result = json.loads(response_body.decode('utf-8'))

        if "items" in result and len(result["items"]) > 0:
            description_all = ' '.join([item["description"] for item in result["items"]])
            clean_description = re.sub(r'<\/?b>', '', description_all)
            print(f"참고자료: {clean_description}")
        else:
            clean_description = "참고 자료가 없습니다."
    except Exception as e:
        print("블로그 검색 중 오류가 발생했습니다.")
        print(str(e))
        clean_description = "참고 자료가 없습니다."
    return clean_description

def read_user_example_text(file_path):
    # 사용자 예시 텍스트 파일이 존재하는지 확인
    if os.path.exists(file_path):
        with open(file_path, 'r', encoding='utf-8') as file:
            example_text = file.read()
        print("사용자 예시 텍스트를 성공적으로 로드했습니다.")
    else:
        example_text = ""
        print("user_example_text.txt 파일을 찾을 수 없어 예시 텍스트 없이 진행합니다.")
    return example_text

def choose_tone(tone_choice):
    # 톤 선택 사전
    tones = {
        "1": "formal",      # 공식적/비즈니스 어조
        "2": "casual",      # 친근하고 가벼운 어조
        "3": "humorous",    # 유머러스한 어조
        "4": "informative"  # 정보 제공형 어조
    }
    # 선택한 톤 반환, 기본은 'casual'
    return tones.get(str(tone_choice), "casual")

def generate_final_post(second_sys_prompt_content, third_sys_prompt_content, user_question, clean_description, image_captions, example_text, tone=None):
    # 톤 프롬프트에 추가 (tone이 있으면 해당하는 프롬프트 추가)
    tone_instruction = f"Please write in a {tone} tone." if tone else ""
    example_text_content = f"Here is an example of the user's previous writing style: {example_text}" if example_text else "The user has not provided an example text."

    final_completion = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": f"{second_sys_prompt_content}\n\n{tone_instruction}\n\n글 형식:{third_sys_prompt_content}"
            },
            {
                "role": "user",
                "content": f"사용자의 질문: {user_question}\n참고자료: {clean_description}\n입력된 사진: {' '.join(image_captions)}\n{example_text_content}"
            }
        ]
    )

    final_post = final_completion.choices[0].message.content.strip()
    return final_post

def apply_md_formatting(paragraph, text):
    # 기존 내용 제거
    p_element = paragraph._element
    for child in p_element[:]:
        p_element.remove(child)

    # 헤딩 처리
    heading_match = re.match(r'^(#{1,6})\s+(.*)', text)
    if heading_match:
        level = len(heading_match.group(1))
        content = heading_match.group(2)
        paragraph.style = f'Heading {level}'
        process_inline_formatting(paragraph, content)
        return

    # 리스트 처리
    list_match = re.match(r'^(\*|\-|\+)\s+(.*)', text)
    if list_match:
        content = list_match.group(2)
        paragraph.style = 'List Bullet'
        process_inline_formatting(paragraph, content)
        return

    numbered_list_match = re.match(r'^(\d+)\.\s+(.*)', text)
    if numbered_list_match:
        content = numbered_list_match.group(2)
        paragraph.style = 'List Number'
        process_inline_formatting(paragraph, content)
        return

    # 블록 인용 처리
    blockquote_match = re.match(r'^>\s+(.*)', text)
    if blockquote_match:
        content = blockquote_match.group(1)
        paragraph.style = 'Quote'
        process_inline_formatting(paragraph, content)
        return

    # 수평선 처리
    if re.match(r'^(\*\*\*|---)$', text.strip()):
        run = paragraph.add_run()
        p = run._element
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.insert_element_before(pBdr, 'w:spacing')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        return

    # 코드 블록 처리
    code_block_match = re.match(r'^```(.*)', text)
    if code_block_match:
        content = code_block_match.group(1)
        run = paragraph.add_run(content)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        return

    # 인라인 서식 처리
    process_inline_formatting(paragraph, text)

def process_inline_formatting(paragraph, text):
    # 인라인 마크다운 서식을 적용합니다.
    # 패턴 정의
    pattern = r'(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*|`.+?`|~~.+?~~|\!\[.*?\]\(.*?\)|\[.+?\]\(.*?\))'
    tokens = re.split(pattern, text)

    for token in tokens:
        if not token:
            continue
        if token.startswith('***') and token.endswith('***'):
            # 굵게 및 기울임
            run = paragraph.add_run(token[3:-3])
            run.bold = True
            run.italic = True
        elif token.startswith('**') and token.endswith('**'):
            # 굵게
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            # 기울임
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith('~~') and token.endswith('~~'):
            # 취소선
            run = paragraph.add_run(token[2:-2])
            run.font.strike = True
        elif token.startswith('`') and token.endswith('`'):
            # 인라인 코드
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        elif token.startswith('![') and '](' in token and token.endswith(')'):
            # 마크다운 이미지 태그 처리 (![alt](image.png))
            alt_text = token[2:token.index('](')]
            image_url = token[token.index('](')+2:-1]
            # 현재 Word 문서에 이미지를 직접 삽입하므로 텍스트로 변환
            run = paragraph.add_run(f"[이미지: {alt_text}]")
            run.font.color.rgb = RGBColor(0, 0, 255)  # 파란색으로 표시
            run.font.underline = True
        elif token.startswith('[') and '](' in token and token.endswith(')'):
            # 하이퍼링크
            link_text = token[1:token.index(']')]
            link_url = token[token.index('](')+2:-1]
            run = paragraph.add_run(f'{link_text} ({link_url})')
            run.font.color.rgb = RGBColor(0, 0, 255)  # 파란색으로 표시
            run.font.underline = True
        else:
            # 일반 텍스트
            run = paragraph.add_run(token)

def save_post_to_word(final_post):
    # Word 파일에 저장할 폴더 경로 설정
    output_folder = "/output/"
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    doc = Document()
    
    # 이미지 태그 패턴 정의
    image_tag_patterns = [
        r'\{\{(.+?\.(?:png))\}\}',  # {{image.png}}
        r'\((.+?\.(?:png))\)',      # (image.png)
        r'!\[.*?\]\((.+?\.(?:png))\)',  # ![alt](image.png)
        r'!\[(.*?)\]'               # ![alt]
    ]
    
    # 문서에 텍스트를 삽입하면서 이미지 태그와 마크다운 서식을 인식
    for line in final_post.split('\n'):
        # 모든 이미지 태그 패턴에 대해 반복
        for pattern in image_tag_patterns:
            matches = re.findall(pattern, line)
            for image_name in matches:
                if pattern == r'!\[(.*?)\]':
                    # ![alt] 형식의 경우 이미지 이름을 alt 텍스트로 가정
                    image_name = image_name + ".png"  # 확장자 추가 필요 시 조정

                image_path = os.path.join("/data/test/", image_name)
                
                # 이미지가 존재하는 경우 이미지 삽입
                if os.path.exists(image_path):
                    try:
                        doc.add_picture(image_path, width=Inches(5))  # 이미지 삽입 (너비 5인치로 설정)
                        print(f"이미지 삽입 완료: {image_name}")
                    except Exception as e:
                        print(f"이미지 삽입 실패: {image_name}, 오류: {e}")
                        # 이미지 삽입 실패 시 대체 텍스트 삽입
                        line = re.sub(rf'!\[.*?\]\({re.escape(image_name)}\)', f"[이미지 '{image_name}'를 삽입할 수 없습니다]", line)
                
                # 이미지 태그를 제거
                if pattern == r'\{\{(.+?\.(?:png))\}\}':
                    line = re.sub(r'\{\{'+re.escape(image_name)+r'\}\}', '', line)
                elif pattern == r'\((.+?\.(?:png))\)':
                    line = re.sub(r'\('+re.escape(image_name)+r'\)', '', line)
                elif pattern == r'!\[.*?\]\((.+?\.(?:png))\)':
                    line = re.sub(r'!\[.*?\]\('+re.escape(image_name)+r'\)', '', line)
                elif pattern == r'!\[(.*?)\]':
                    line = re.sub(r'!\['+re.escape(image_name[:-4])+r'\]', '', line)
        
        # 이미지 태그가 제거된 후의 텍스트 삽입
        if line.strip():  # 빈 줄이 아닌 경우
            paragraph = doc.add_paragraph()
            apply_md_formatting(paragraph, line)
    
    # 파일 이름 설정 및 저장
    output_file = os.path.join(output_folder, "generated_post_with_images.docx")
    doc.save(output_file)
    
    print(f"게시글이 {output_file} 파일로 저장되었습니다.")

def main():
    folder_path = "/data/test/"
    # 모든 이미지를 PNG로 변환
    convert_images_to_png(folder_path)
    
    api_key, client_id, client_secret = get_api_keys()
    create_openai_client(api_key)
    
    first_sys_prompt = read_sys_prompt('/data/1st_sys_prompt.json')
    user_question = "애플 맥북 m2과 m3의 성능비교에 대한 게시글 작성해줘."
    keyword = generate_keywords(first_sys_prompt["content"], user_question)
    
    clean_description = search_naver_blog(client_id, client_secret, keyword)
    second_sys_prompt = read_sys_prompt('/data/2nd_sys_prompt.json')
    third_sys_prompt = read_sys_prompt('/data/3rd_sys_prompt.json')
    
    # 사용자 예시 텍스트 읽기 (없으면 빈 문자열로 처리)
    example_text = read_user_example_text('/data/user_example_text.txt')
    
    # 원하는 글 형식을 선택하여 사용 (예: 'instagram', 'naver_blog' 등)
    chosen_format = third_sys_prompt["formats"]["naver_blog"]
    
    # 톤 선택: 번호에 따라 톤 결정
    tone_choice = 1  # 1: formal, 2: casual, 3: humorous, 4: informative
    tone = choose_tone(tone_choice)

    # 이미지 분석 (PNG 이미지만 처리)
    image_captions, image_filenames = analyze_images_in_folder()
    final_post = generate_final_post(second_sys_prompt["content"], chosen_format, user_question, clean_description, image_captions, example_text, tone)
    
    save_post_to_word(final_post)

if __name__ == "__main__":
    main()





    # 메인 헤더 스타일링 (글자 크기 조절 가능)
    header_font_size = 36  # 원하는 글자 크기로 변경하세요
    st.markdown(
        f"""
        <h1 style='text-align: center; color: #4b8bbe; font-size: {header_font_size}px;'>📄 이미지 및 텍스트 분석을 통한 자동 게시글 생성 시스템</h1>
        """,
        unsafe_allow_html=True,
    )