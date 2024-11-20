# src/app.py
import os
import re
import json
import urllib.request
import urllib.parse
from openai import OpenAI
import streamlit as st
from PIL import Image
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from transformers import BlipProcessor, BlipForConditionalGeneration
import base64

# Streamlit 페이지 기본 설정
st.set_page_config(
    page_title="자동 게시글 생성 시스템",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded",
)

def add_custom_css():
    # 배경 이미지 로드 및 base64 인코딩
    try:
        with open('data/main_image.png', 'rb') as image_file:
            encoded_string = base64.b64encode(image_file.read()).decode()
    except FileNotFoundError:
        encoded_string = ''  # 배경 이미지가 없을 경우 빈 문자열로 설정

    st.markdown(
        f"""
        <style>
        /* 전체 배경 이미지 설정 */
        body {{
            background-image: url(data:image/png;base64,{encoded_string});
            background-size: cover;
            background-repeat: no-repeat;
            background-attachment: fixed;
            background-position: center;
        }}

        /* 메인 컨테이너 배경 투명하게 */
        .main .block-container {{
            background-color: rgba(255, 255, 255, 0.8);
            padding: 20px;
            border-radius: 10px;
        }}

        /* 헤더 스타일 */
        .main-header {{
            text-align: center;
            color: #2c3e50;
            margin-top: 20px;
            margin-bottom: 40px;
        }}

        /* 입력 섹션 스타일 */
        .stSidebar {{
            background-color: #ecf0f1;
        }}

        /* 진행 과정 박스 스타일 */
        .progress-box {{
            background-color: rgba(255, 255, 255, 0.9);
            padding: 20px;
            border-radius: 10px;
            margin-bottom: 20px;
        }}

        /* 생성된 게시글 박스 스타일 */
        .generated-post {{
            background-color: rgba(255, 255, 255, 0.9);
            padding: 30px;
            border-radius: 10px;
            margin-bottom: 20px;
        }}

        /* 푸터 스타일 */
        .footer {{
            position: fixed;
            right: 10px;
            bottom: 10px;
            text-align: right;
            font-size: 12px;
            color: #95a5a6;
        }}

        /* 반응형 디자인 */
        @media (max-width: 768px) {{
            .main-header h1 {{
                font-size: 24px;
            }}
            .generated-post {{
                padding: 20px;
            }}
        }}
        </style>
        """,
        unsafe_allow_html=True,
    )

def get_api_keys():
    # secrets.toml 파일에서 정보 가져오기 또는 환경 변수에서 가져오기
    api_key = st.secrets.get("OPENAI_API_KEY") or os.getenv("OPENAI_API_KEY")
    client_id = st.secrets.get("NAVER_CLIENT_ID") or os.getenv("NAVER_CLIENT_ID")
    client_secret = st.secrets.get("NAVER_CLIENT_SECRET") or os.getenv("NAVER_CLIENT_SECRET")
    
    if not api_key or not client_id or not client_secret:
        st.error("OPENAI_API_KEY, NAVER_CLIENT_ID 또는 NAVER_CLIENT_SECRET이 설정되지 않았습니다.")
        st.stop()
    
    return api_key, client_id, client_secret

def create_openai_client(api_key):
    # OpenAI 클라이언트 초기화
    client = OpenAI(api_key=api_key)
    return client

def analyze_uploaded_images(uploaded_images, progress_messages):
    # BLIP 프로세서 및 모델 초기화
    processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base")
    model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base")
    
    image_filenames = []   # 이미지 파일명만 저장
    image_bytes_dict = {}  # 이미지 파일명과 바이트 데이터를 저장
    
    for uploaded_file in uploaded_images:
        image = Image.open(uploaded_file)
        if image.mode != 'RGBA':
            image = image.convert('RGBA')
        inputs = processor(images=image, return_tensors="pt")
        out = model.generate(**inputs, max_new_tokens=50)
        caption = processor.decode(out[0], skip_special_tokens=True)
        image_filenames.append(uploaded_file.name)
        
        # 이미지 바이트 읽기
        uploaded_file.seek(0)  # 파일 포인터를 처음으로 이동
        image_bytes = uploaded_file.read()
        image_bytes_dict[uploaded_file.name] = image_bytes
        
        progress_messages.append(f"파일: {uploaded_file.name}, 이미지 설명: {caption}")
    
    return image_filenames, image_bytes_dict


def read_sys_prompt(prompt_name):
    # 기본 프롬프트 설정
    prompts = {
        "first_sys_prompt": {
            "content": "당신은 키워드 추출 전문가입니다. 사용자의 질문에서 핵심 키워드를 추출하세요."
        },
        "second_sys_prompt": {
            "content": "당신은 글 작성 전문가입니다. 아래의 정보를 바탕으로 게시글을 작성하세요.\n이미지를 글에 포함시킬 때는 이미지 파일명을 중괄호로 감싸서 {image_filename} 형태로 표시하세요."
        }
    }
    return prompts.get(prompt_name, {"content": ""})

def load_third_sys_prompt():
    # '3rd_sys_prompt.json' 파일 로드
    try:
        with open('data/3rd_sys_prompt.json', 'r', encoding='utf-8') as f:
            third_sys_prompt = json.load(f)
        return third_sys_prompt
    except Exception as e:
        st.error("3rd_sys_prompt.json 파일을 로드하는 데 실패했습니다.")
        st.error(str(e))
        st.stop()

def generate_keywords(client, first_sys_prompt_content, user_question, language):
    completion = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": f"{first_sys_prompt_content}\n사용 언어: {language}"
            },
            {
                "role": "user",
                "content": user_question
            }
        ]
    )

    keyword = completion.choices[0].message.content.strip()
    return keyword

def search_naver_blog(client_id, client_secret, keyword):
    encText = urllib.parse.quote(keyword)
    display = 10
    start = 1
    sort = "sim"
    
    url = f"https://openapi.naver.com/v1/search/blog?query={encText}&display={display}&start={start}&sort={sort}"
    
    request = urllib.request.Request(url)
    request.add_header("X-Naver-Client-Id", client_id)
    request.add_header("X-Naver-Client-Secret", client_secret)
    
    try:
        response = urllib.request.urlopen(request)
        response_body = response.read()
        result = json.loads(response_body.decode('utf-8'))

        if "items" in result and len(result["items"]) > 0:
            description_all = ' '.join([item["description"] for item in result["items"]])
            clean_description = re.sub(r'<\/?b>', '', description_all)
        else:
            clean_description = "참고 자료가 없습니다."
    except Exception as e:
        clean_description = "참고 자료가 없습니다."
    return clean_description

def choose_tone(tone_choice):
    # 톤 선택 사전
    tones = {
        "1": "formal",      # 공식적/비즈니스 어조
        "2": "casual",      # 친근하고 가벼운 어조
        "3": "humorous",    # 유머러스한 어조
        "4": "informative"  # 정보 제공형 어조
    }
    # 선택한 톤 반환, 기본은 'casual'
    return tones.get(str(tone_choice), "casual")

def generate_final_post(client, second_sys_prompt_content, chosen_format_content, user_question, clean_description, image_filenames, example_text, tone, language):
    # 톤 프롬프트에 추가 (tone이 있으면 해당하는 프롬프트 추가)
    tone_instruction = f"Please write in a {tone} tone." if tone else ""
    example_text_content = f"Here is an example of the user's previous writing style: {example_text}" if example_text else "The user has not provided an example text."

    # 이미지 삽입에 대한 지시 추가
    image_instructions = "Please include the images in the generated text at appropriate positions using the format {image_filename}."

    # 이미지 파일명을 {ham1.jpeg} 형태로 포맷
    image_placeholders = ' '.join([f'{{{fn}}}' for fn in image_filenames])

    final_completion = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": f"{second_sys_prompt_content}\n\n{tone_instruction}\n\n{image_instructions}\n\n글 형식: {chosen_format_content}\n사용 언어: {language}"
            },
            {
                "role": "user",
                "content": f"사용자의 질문: {user_question}\n참고자료: {clean_description}\n입력된 사진: {image_placeholders}\n{example_text_content}"
            }
        ]
    )

    final_post = final_completion.choices[0].message.content.strip()
    return final_post

def translate_post(client, final_post, target_language):
    # 게시글 번역 함수 추가
    translation_completion = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": f"Please translate the following text into {target_language}. Maintain the formatting and placeholders for images (e.g., {{image_filename}})."
            },
            {
                "role": "user",
                "content": final_post
            }
        ]
    )
    translated_post = translation_completion.choices[0].message.content.strip()
    return translated_post

def apply_md_formatting(paragraph, text):
    # 기존 내용 제거
    p_element = paragraph._element
    for child in p_element[:]:
        p_element.remove(child)

    # 헤딩 처리
    heading_match = re.match(r'^(#{1,6})\s+(.*)', text)
    if heading_match:
        level = len(heading_match.group(1))
        content = heading_match.group(2)
        paragraph.style = f'Heading {level}'
        # 중앙 정렬
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        process_inline_formatting(paragraph, content)
        return

    # 리스트 처리
    list_match = re.match(r'^(\*|\-|\+)\s+(.*)', text)
    if list_match:
        content = list_match.group(2)
        paragraph.style = 'List Bullet'
        process_inline_formatting(paragraph, content)
        return

    numbered_list_match = re.match(r'^(\d+)\.\s+(.*)', text)
    if numbered_list_match:
        content = numbered_list_match.group(2)
        paragraph.style = 'List Number'
        process_inline_formatting(paragraph, content)
        return

    # 블록 인용 처리
    blockquote_match = re.match(r'^>\s+(.*)', text)
    if blockquote_match:
        content = blockquote_match.group(1)
        paragraph.style = 'Quote'
        process_inline_formatting(paragraph, content)
        return

    # 수평선 처리
    if re.match(r'^(\*\*\*|---)$', text.strip()):
        run = paragraph.add_run()
        p = run._element
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.insert_element_before(pBdr, 'w:spacing')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        return

    # 코드 블록 처리
    code_block_match = re.match(r'^```(.*)', text)
    if code_block_match:
        content = code_block_match.group(1)
        run = paragraph.add_run(content)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        return

    # 인라인 서식 처리
    process_inline_formatting(paragraph, text)

def process_inline_formatting(paragraph, text):
    # 인라인 마크다운 서식을 적용합니다.
    # 패턴 정의
    pattern = r'(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*|`.+?`|~~.+?~~|\!\[.*?\]\(.*?\)|\!\[.*?\]|\[.+?\]\(.*?\))'
    tokens = re.split(pattern, text)

    for token in tokens:
        if not token:
            continue
        if token.startswith('***') and token.endswith('***'):
            # 굵게 및 기울임
            run = paragraph.add_run(token[3:-3])
            run.bold = True
            run.italic = True
        elif token.startswith('**') and token.endswith('**'):
            # 굵게
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            # 기울임
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith('~~') and token.endswith('~~'):
            # 취소선
            run = paragraph.add_run(token[2:-2])
            run.font.strike = True
        elif token.startswith('`') and token.endswith('`'):
            # 인라인 코드
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        elif token.startswith('![') and '](' in token and token.endswith(')'):
            # 마크다운 이미지 태그 처리 (![alt](image.png))
            alt_text = token[2:token.index('](')]
            image_url = token[token.index('](')+2:-1]
            # 텍스트로 대체
            run = paragraph.add_run(f"[이미지: {alt_text}]")
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.font.underline = True
        elif token.startswith('![') and token.endswith(']'):
            # 마크다운 이미지 태그 처리 (![alt])
            alt_text = token[2:-1]
            run = paragraph.add_run(f"[이미지: {alt_text}]")
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.font.underline = True
        elif token.startswith('[') and '](' in token and token.endswith(')'):
            # 하이퍼링크
            link_text = token[1:token.index(']')]
            link_url = token[token.index('](')+2:-1]
            run = paragraph.add_run(f'{link_text} ({link_url})')
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.font.underline = True
        else:
            # 일반 텍스트
            run = paragraph.add_run(token)

def save_post_to_word(final_post, image_bytes_dict):
    doc = Document()

    # 이미지 태그 패턴 정의
    image_tag_patterns = [
        r'\{(.+?\.(?:png|jpg|jpeg))\}',       # {image.png}
        r'\((.+?\.(?:png|jpg|jpeg))\)',      # (image.png)
        r'!\[.*?\]\((.+?\.(?:png|jpg|jpeg))\)',  # ![alt](image.png)
        r'!\[(.*?)\]'                          # ![alt]
    ]

    for line in final_post.split('\n'):
        original_line = line  # 디버깅용 원본 라인 저장

        # 이미지 태그 처리
        for pattern in image_tag_patterns:
            matches = re.findall(pattern, line)
            for image_name in matches:
                if pattern == r'!\[(.*?)\]':
                    # ![alt] 형식의 경우 이미지 이름을 alt 텍스트로 가정
                    image_name = image_name + ".png"  # 확장자 추가 필요 시 조정

                image_bytes = image_bytes_dict.get(image_name)
                if image_bytes:
                    try:
                        # 이미지 바이트를 BytesIO로 변환
                        image_stream = BytesIO(image_bytes)
                        doc.add_picture(image_stream, width=Inches(5))
                    except Exception as e:
                        line = line.replace(image_name, f"[이미지 '{image_name}'를 삽입할 수 없습니다]")
                else:
                    line = line.replace(image_name, f"[이미지 '{image_name}'를 찾을 수 없습니다]")

                # 이미지 태그 제거
                if pattern == r'\{(.+?\.(?:png|jpg|jpeg))\}':
                    line = re.sub(r'\{' + re.escape(image_name) + r'\}', '', line)
                elif pattern == r'\((.+?\.(?:png|jpg|jpeg))\)':
                    line = re.sub(r'\(' + re.escape(image_name) + r'\)', '', line)
                elif pattern == r'!\[.*?\]\((.+?\.(?:png|jpg|jpeg))\)':
                    line = re.sub(r'!\[.*?\]\(' + re.escape(image_name) + r'\)', '', line)
                elif pattern == r'!\[(.*?)\]':
                    line = re.sub(r'!\[' + re.escape(image_name[:-4]) + r'\]', '', line)

        # 이미지 태그가 제거된 후의 텍스트 삽입
        if line.strip():
            paragraph = doc.add_paragraph()
            apply_md_formatting(paragraph, line)

    # BytesIO를 사용하여 메모리에 저장
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def main():
    # 세션 상태 초기화
    if 'generated_post' not in st.session_state:
        st.session_state['generated_post'] = ''
    if 'progress_messages' not in st.session_state:
        st.session_state['progress_messages'] = []
    if 'translated_posts' not in st.session_state:
        st.session_state['translated_posts'] = {}
    if 'image_bytes_dict' not in st.session_state:
        st.session_state['image_bytes_dict'] = {}

    # 커스텀 CSS 추가
    add_custom_css()

    # 메인 헤더 스타일링 (글자 크기 조절 가능)
    header_font_size = 36  # 원하는 글자 크기로 변경하세요
    st.markdown(
        f"""
        <h1 style='text-align: center; color: #4b8bbe; font-size: {header_font_size}px;'>📄 이미지 및 텍스트 분석을 통한 자동 게시글 생성 시스템</h1>
        """,
        unsafe_allow_html=True,
    )

    # API 키 가져오기
    api_key, client_id, client_secret = get_api_keys()
    client = create_openai_client(api_key)  # OpenAI 클라이언트 초기화

    # 시스템 프롬프트 읽기
    first_sys_prompt = read_sys_prompt('first_sys_prompt')
    second_sys_prompt = read_sys_prompt('second_sys_prompt')
    third_sys_prompt = load_third_sys_prompt()

    # 사이드바 입력
    st.sidebar.header("📝 입력 설정")
    user_question = st.sidebar.text_area("작성하고자 하는 내용 입력", "")

    # 언어 선택 수정 (다중 선택 가능하도록 변경)
    languages = ["한국어", "English", "日本語", "中文", "Español", "Français"]
    language_choices = st.sidebar.multiselect("언어 선택 (여러 개 선택 가능)", options=languages, default=["한국어"])

    # 톤 선택
    tone_options = {
        "1": "공식적/비즈니스",  # formal
        "2": "친근하고 가벼운",  # casual
        "3": "유머러스한",       # humorous
        "4": "정보 제공형"       # informative
    }
    tone_choice = st.sidebar.selectbox("톤 선택", options=list(tone_options.keys()), format_func=lambda x: tone_options[x])
    tone = choose_tone(tone_choice)

    # 글 형식 선택
    format_options = list(third_sys_prompt["formats"].keys())
    format_choice = st.sidebar.selectbox("글 형식 선택", options=format_options, format_func=lambda x: x)
    chosen_format_content = third_sys_prompt["formats"][format_choice]

    # 이미지 업로드
    uploaded_images = st.sidebar.file_uploader("이미지 업로드", accept_multiple_files=True, type=["jpg", "jpeg", "png"])

    # 예시 텍스트 입력 또는 파일 업로드
    st.sidebar.write("예시 텍스트 입력 또는 파일 업로드 (선택사항)")
    example_text = st.sidebar.text_area("예시 텍스트 입력", "")
    example_file = st.sidebar.file_uploader("예시 텍스트 파일 업로드 (.txt)", type=["txt"])

    if example_file is not None:
        example_text = example_file.read().decode('utf-8')
        st.sidebar.success("예시 텍스트 파일이 업로드되었습니다.")

    if st.sidebar.button("📄 게시글 생성"):
        st.session_state['progress_messages'] = []  # 진행 과정 초기화
        st.session_state['translated_posts'] = {}    # 번역된 게시글 초기화

        if not user_question:
            st.error("작성하고자 하는 내용을 입력하세요.")
            return

        # 이미지 캡션 및 바이트 생성
        if uploaded_images:
            with st.spinner("이미지 분석 중..."):
                image_filenames, image_bytes_dict = analyze_uploaded_images(uploaded_images, st.session_state['progress_messages'])
        else:
            image_filenames = []
            image_bytes_dict = {}
            st.session_state['progress_messages'].append("이미지가 업로드되지 않았습니다.")

        # 키워드 생성 (첫 번째 선택한 언어로)
        if language_choices:
            with st.spinner("키워드 생성 중..."):
                keyword = generate_keywords(client, first_sys_prompt["content"], user_question, language_choices[0])
                st.session_state['progress_messages'].append(f"추출된 키워드: {keyword}")
        else:
            st.error("언어를 선택하세요.")
            return

        # 네이버 블로그 검색
        with st.spinner("블로그에서 참고자료 수집 중..."):
            clean_description = search_naver_blog(client_id, client_secret, keyword)
            st.session_state['progress_messages'].append("참고자료 수집 완료")

        # 게시글 생성 (첫 번째 선택한 언어로)
        with st.spinner("게시글 생성 중..."):
            final_post = generate_final_post(client, second_sys_prompt["content"], chosen_format_content, user_question, clean_description, image_filenames, example_text, tone, language_choices[0])
            st.session_state['generated_post'] = final_post  # 세션 상태에 저장

        # 이미지 바이트를 세션 상태에 저장
        st.session_state['image_bytes_dict'] = image_bytes_dict

        # 선택된 다른 언어로 번역
        if len(language_choices) > 1:
            for lang in language_choices[1:]:
                with st.spinner(f"{lang}로 번역 중..."):
                    translated_post = translate_post(client, final_post, lang)
                    st.session_state['translated_posts'][lang] = translated_post
                    st.session_state['progress_messages'].append(f"{lang}로 번역 완료")

    # 진행 과정 표시
    if st.session_state['progress_messages']:
        st.markdown("## 🔄 진행 과정")
        st.markdown('<div class="progress-box">', unsafe_allow_html=True)
        for msg in st.session_state['progress_messages']:
            st.write(f"- {msg}")
        st.markdown('</div>', unsafe_allow_html=True)

    # 생성된 게시글 표시
    if st.session_state['generated_post']:
        st.markdown("## ✨ 생성된 게시글")
        st.markdown('<div class="generated-post">', unsafe_allow_html=True)
        st.markdown(f"### {language_choices[0]}")
        st.markdown(st.session_state['generated_post'], unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

        # 게시글 Word 파일로 저장
        word_file = save_post_to_word(st.session_state['generated_post'], st.session_state.get('image_bytes_dict', {}))
        st.download_button(
            label=f"📥 게시글 Word 파일로 다운로드 ({language_choices[0]})",
            data=word_file,
            file_name=f"generated_post_{language_choices[0]}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # 번역된 게시글 표시 및 다운로드
        for lang, translated_post in st.session_state['translated_posts'].items():
            st.markdown('<div class="generated-post">', unsafe_allow_html=True)
            st.markdown(f"### {lang}")
            st.markdown(translated_post, unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)

            # 번역된 게시글 Word 파일로 저장
            translated_word_file = save_post_to_word(translated_post, st.session_state.get('image_bytes_dict', {}))
            st.download_button(
                label=f"📥 게시글 Word 파일로 다운로드 ({lang})",
                data=translated_word_file,
                file_name=f"generated_post_{lang}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    # 푸터 추가
    st.markdown(
        """
        <div class="footer">
            Made by Team Nice
        </div>
        """,
        unsafe_allow_html=True,
    )

if __name__ == "__main__":
    main()
